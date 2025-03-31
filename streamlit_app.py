import streamlit as st
import os
import fitz as pymupdf
from docx import Document
from dotenv import load_dotenv
import google.generativeai as genai
from io import BytesIO
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import zipfile  # Add this import at the top of your file

# Import the prompt function from the prompts directory
from prompts.lesson_blueprint_prompt import get_prompt
from prompts.assessment_items_prompt import get_prompt as get_assessment_prompt
from prompts.media_suggestions_prompt import get_prompt as get_media_prompt
from prompts.fact_check_prompt import get_prompt as get_fact_check_prompt
from prompts.dei_check_prompt import get_prompt as get_dei_check_prompt


# Load environment variables
load_dotenv()
genai_api_key = os.getenv("GEMINI_API_KEY")

if not genai_api_key:
    st.error("GEMINI_API_KEY not found. Please check your .env file.")
    st.stop()

genai.configure(api_key=genai_api_key)
# Model instances
blueprint_model = genai.GenerativeModel("gemini-1.5-flash")
assessment_model = genai.GenerativeModel("gemini-1.5-flash")
media_model = genai.GenerativeModel("gemini-1.5-flash")  
fact_check_model = genai.GenerativeModel("gemini-1.5-flash")  
dei_check_model = genai.GenerativeModel("gemini-1.5-flash") 


# Initialize session state variables if they don't exist
if 'blueprint_output' not in st.session_state:
    st.session_state.blueprint_output = None
if 'assessment_output' not in st.session_state:
    st.session_state.assessment_output = None
if 'media_output' not in st.session_state:
    st.session_state.media_output = None
if 'fact_check_output' not in st.session_state:
    st.session_state.fact_check_output = None
if 'dei_check_output' not in st.session_state:
    st.session_state.dei_check_output = None
if 'has_generated' not in st.session_state:
    st.session_state.has_generated = False

# Add reset function
def reset_outputs():
    st.session_state.blueprint_output = None
    st.session_state.assessment_output = None
    st.session_state.media_output = None
    st.session_state.fact_check_output = None
    st.session_state.dei_check_output = None
    st.session_state.has_generated = False


# Helper functions
def extract_text_from_pdf(pdf_path):
    try:
        doc = pymupdf.open(pdf_path)
        text = ''
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        return f"Error reading PDF ({pdf_path}): {str(e)}"

def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error reading DOCX ({docx_path}): {str(e)}"

def load_reference_materials(folder_path):
    reference_texts = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif filename.lower().endswith(('.txt', '.md')):
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        elif filename.lower().endswith('.docx'):
            text = extract_text_from_docx(file_path)
        else:
            continue
        reference_texts.append(f"Document: {filename}\n{text}\n\n")
    return "\n".join(reference_texts)

# Updated function to load specific reference materials (PDF and DOCX)
def load_dei_reference_materials(folder_path):
    dei_reference_texts = []
    target_files = ["Subject Specific Guidelines-Social Studies", "DEI Content Authoring Guidelines"]
    
    for filename in os.listdir(folder_path):
        # Check if the filename contains any of the target file names
        if any(target_file in filename for target_file in target_files):
            file_path = os.path.join(folder_path, filename)
            
            # Process based on file extension
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.lower().endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                continue
                
            dei_reference_texts.append(f"Document: {filename}\n{text}\n\n")
    
    return "\n".join(dei_reference_texts)

# Function to load specific blueprint reference material
def load_blueprint_reference(folder_path):
    target_file = "CCAG-EdgeEX Lesson Blueprinting-270325-194434"
    
    for filename in os.listdir(folder_path):
        if target_file in filename:
            file_path = os.path.join(folder_path, filename)
            if filename.lower().endswith('.pdf'):
                return extract_text_from_pdf(file_path)
            elif filename.lower().endswith('.docx'):
                return extract_text_from_docx(file_path)
    
    # If file not found, return an empty string or a message
    return "Blueprint reference file not found."


# Load reference materials
reference_materials_folder = 'reference_materials'
reference_content = load_reference_materials(reference_materials_folder)

# Load DEI-specific reference materials
dei_reference_content = load_dei_reference_materials(reference_materials_folder)

# Load blueprint-specific reference material
blueprint_reference_content = load_blueprint_reference(reference_materials_folder)

# Streamlit page setup
st.set_page_config(page_title="Lesson Blueprint Generator", layout="centered")
st.title("Lesson Blueprint, Assessment, and Media Suggestion Generator")

# Input fields
lesson_title = st.text_input("Lesson Title", key="lesson_title")
lesson_info = st.text_area("Lesson Information", 
                          "Enter lesson description, question, and learning objectives here.", 
                          height=200)
additional_resources = st.text_area("Additional Resources", 
                                  "Enter any additional resources or notes here.", 
                                  height=100)


# Instruction prompt for Lesson Blueprint Generation
def create_lesson_blueprint(lesson_title, lesson_info, additional_resources, blueprint_reference_content):
    # Get the prompt from the imported function, using the specific blueprint reference
    prompt = get_prompt(blueprint_reference_content, lesson_info, additional_resources, lesson_title)
    
    # Add DEI reference materials to consider
    prompt += f"""
    
    ### Additional DEI Considerations
    Please ensure your lesson blueprint incorporates these DEI guidelines:
    {dei_reference_content}
    """
    
    try:
        response = blueprint_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating content: {str(e)}"

# Instruction prompt for Assessment Items Generation
def create_assessment_items(blueprint):
    prompt = get_assessment_prompt(blueprint)
    
    # Add DEI reference materials to consider
    prompt += f"""
    
    ### Additional DEI Considerations
    Please ensure your assessment items follow these DEI guidelines:
    {dei_reference_content}
    """
    
    try:
        response = assessment_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating assessment items: {str(e)}"

# Instruction prompt for Media Suggestions Generation
def create_media_suggestions(blueprint, assessment):
    prompt = get_media_prompt(blueprint, assessment)
    
    # Add DEI reference materials to consider
    prompt += f"""
    
    ### Additional DEI Considerations
    Please ensure your media suggestions follow these DEI guidelines:
    {dei_reference_content}
    """
    
    try:
        response = media_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating media suggestions: {str(e)}"
    
# Functions for fact checking and DEI checking
def create_fact_check(blueprint, assessment, media_suggestions):
    prompt = get_fact_check_prompt(blueprint, assessment, media_suggestions)
    
    # Add DEI reference materials for context
    prompt += f"""
    
    ### Additional Context
    While conducting the fact check, please be aware of these DEI guidelines:
    {dei_reference_content}
    """
    
    try:
        response = fact_check_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating fact check: {str(e)}"

def create_dei_check(blueprint, assessment, media_suggestions):
    prompt = get_dei_check_prompt(blueprint, assessment, media_suggestions, dei_reference_content)
    try:
        response = dei_check_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating DEI check: {str(e)}"

# Function to create a Word document
def create_word_doc(title, content):
    doc = Document()
    # Add title with formatting
    title_heading = doc.add_heading(title, level=1)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Split content into paragraphs
    paragraphs = content.split('\n')
    
    # Flag to track if we're inside a table
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        if not para.strip():
            i += 1
            continue
        
        # Check for table start (line contains | and the next line has dashes and |)
        if '|' in para and i + 1 < len(paragraphs) and set(paragraphs[i+1].replace('|', '')).issubset({'-', ' '}):
            in_table = True
            table_rows = []
            
            # Add header row
            header_cells = [cell.strip() for cell in para.split('|')[1:-1]] if para.strip().startswith('|') else [cell.strip() for cell in para.split('|')]
            table_rows.append(header_cells)
            
            # Skip the separator row
            i += 2  # Move past header and separator lines
            
            # Process table rows
            while i < len(paragraphs) and '|' in paragraphs[i]:
                row = paragraphs[i]
                row_cells = [cell.strip() for cell in row.split('|')[1:-1]] if row.strip().startswith('|') else [cell.strip() for cell in row.split('|')]
                table_rows.append(row_cells)
                i += 1
            
            # Create the Word table
            if table_rows:
                max_cols = max(len(row) for row in table_rows)
                word_table = doc.add_table(rows=len(table_rows), cols=max_cols)
                word_table.style = 'Table Grid'
                
                # Fill the table
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < max_cols:  # Ensure we don't exceed column count
                            cell = word_table.cell(row_idx, col_idx)
                            cell.text = cell_text
                            
                            # Make header row bold
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                # Add space after table
                doc.add_paragraph()
            
            in_table = False
            continue
            
        # Check if paragraph is a header (markdown style)
        header_match = re.match(r'^#+\s+(.+)$', para)
        if header_match:
            header_text = header_match.group(1)
            level = min(len(re.match(r'^#+', para).group(0)), 6)
            doc.add_heading(header_text, level=level)
            i += 1
            continue
            
        # Check if paragraph is a list item
        if para.strip().startswith('- ') or para.strip().startswith('* '):
            doc.add_paragraph(para.strip()[2:], style='List Bullet')
            i += 1
            continue
            
        # Handle bolded text with markdown ** or __
        bold_pattern = r'\*\*(.*?)\*\*|__(.*?)__'
        if re.search(bold_pattern, para):
            p = doc.add_paragraph()
            parts = re.split(bold_pattern, para)
            for j, part in enumerate(parts):
                if part:  # Skip empty parts
                    if j % 3 == 1 or j % 3 == 2:  # Every third part matches the bold pattern
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            i += 1
            continue
            
        # Regular paragraph
        doc.add_paragraph(para)
        i += 1
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Modify this function to accept existing document buffers instead of recreating them
def create_zip_with_all_docs(blueprint_doc, assessment_doc, media_doc, reports_doc):
    # Create a ZIP file in memory
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
        # Add each document to the zip file
        zip_file.writestr("Lesson_Blueprint.docx", blueprint_doc.getvalue())
        zip_file.writestr("Assessment_Items.docx", assessment_doc.getvalue())
        zip_file.writestr("Media_Suggestions.docx", media_doc.getvalue())
        zip_file.writestr("DEI_Fact_Check_Reports.docx", reports_doc.getvalue())
    
    zip_buffer.seek(0)
    return zip_buffer

# Generate button and workflow
col1, col2 = st.columns([3, 1])
with col1:
    generate_button = st.button("Generate Content")
with col2:
    reset_button = st.button("Reset", on_click=reset_outputs)

# Update your content generation workflow
if generate_button:
    if not lesson_info.strip() or not lesson_title.strip():
        st.warning("Please enter both lesson title and information to generate content.")
    else:
        # Generate primary content
        with st.spinner("Generating lesson blueprint..."):
            st.session_state.blueprint_output = create_lesson_blueprint(
                lesson_title,
                lesson_info, 
                additional_resources, 
                blueprint_reference_content
            )
        
        with st.spinner("Generating assessment items..."):
            st.session_state.assessment_output = create_assessment_items(st.session_state.blueprint_output)
        
        with st.spinner("Generating media suggestions..."):
            st.session_state.media_output = create_media_suggestions(st.session_state.blueprint_output, st.session_state.assessment_output)
        
        # Run fact check and DEI checks in parallel
        col1, col2 = st.columns(2)
        
        with col1:
            with st.spinner("Running fact check..."):
                st.session_state.fact_check_output = create_fact_check(
                    st.session_state.blueprint_output, 
                    st.session_state.assessment_output, 
                    st.session_state.media_output
                )
                
        with col2:
            with st.spinner("Running DEI check..."):
                st.session_state.dei_check_output = create_dei_check(
                    st.session_state.blueprint_output, 
                    st.session_state.assessment_output, 
                    st.session_state.media_output
                )
        
        # Set flag that content has been generated
        st.session_state.has_generated = True
        
        # Use the current rerun method instead of experimental_rerun
        st.rerun()

# Display outputs if they exist
if st.session_state.has_generated:
    # Add a separator before download options
    st.markdown("---")
    st.markdown("### Download Options")

    # Create the four requested document files
    blueprint_doc = create_word_doc("Lesson Blueprint", st.session_state.blueprint_output)
    assessment_doc = create_word_doc("Assessment Items", st.session_state.assessment_output)
    media_doc = create_word_doc("Media Suggestions", st.session_state.media_output)
    reports_doc = create_word_doc(
        "DEI and Fact-check Reports", 
        f"## Fact Check Report\n\n{st.session_state.fact_check_output}\n\n## DEI Check Report\n\n{st.session_state.dei_check_output}"
    )
    
    # Create zip file with all documents - use existing document buffers
    all_docs_zip = create_zip_with_all_docs(
        blueprint_doc,
        assessment_doc,
        media_doc,
        reports_doc
    )

    # Add download buttons in a 5x1 grid for better layout
    col1, col2, col3, col4, col5 = st.columns(5)

    with col1:
        st.download_button(
            label="Download All",
            data=all_docs_zip,
            file_name="All_Lesson_Materials.zip",
            mime="application/zip",
            key="all_docs_download"
        )
    
    with col2:
        st.download_button(
            label="Lesson Blueprint",
            data=blueprint_doc,
            file_name="Lesson_Blueprint.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="blueprint_download"
        )
    with col3:    
        st.download_button(
            label="Assessment Items",
            data=assessment_doc,
            file_name="Assessment_Items.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="assessment_download"
        )
    with col4:
        st.download_button(
            label="Media Suggestions",
            data=media_doc,
            file_name="Media_Suggestions.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="media_download"
        )
    with col5:    
        st.download_button(
            label="Reports",
            data=reports_doc,
            file_name="DEI_Fact_Check_Reports.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="reports_download"
        )
    
    st.markdown("---")
    
    # Display all outputs in expandable sections
    with st.expander("Generated Lesson Blueprint", expanded=True):
        st.write(st.session_state.blueprint_output)
        
    with st.expander("Generated Assessment Items", expanded=True):
        st.write(st.session_state.assessment_output)
        
    with st.expander("Generated Media Suggestions", expanded=True):
        st.write(st.session_state.media_output)
        
    with st.expander("Fact Check Report", expanded=True):
        st.write(st.session_state.fact_check_output)
        
    with st.expander("DEI Check Report", expanded=True):
        st.write(st.session_state.dei_check_output)
